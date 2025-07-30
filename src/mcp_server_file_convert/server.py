import logging
import openpyxl
import pdfplumber
import pandas as pd
from docx import Document
from mcp.server.fastmcp import FastMCP
from pathlib import Path
from typing import Any, Dict

mcp = FastMCP(
    name="file_to_markdown",
    description="将PDF、Word、Excel、csv 转换为Markdown格式的工具",
    version="1.0.0",
)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("file_to_markdown")


async def call_tool(name: str, arguments: Dict[str, Any]) -> str:
    """
    调用指定的工具
    """
    if name == "pdf_to_text":
        return await pdf_to_text(arguments)
    elif name == "docx_to_markdown":
        return await docx_to_markdown(arguments)
    elif name == "excel_to_markdown":
        return await excel_to_markdown(arguments)
    elif name == "csv_to_markdown":
        return await csv_to_markdown(arguments)
    else:
        raise ValueError(f"未知工具: {name}")


async def docx_to_markdown(arguments: Dict[str, Any]) -> str:
    """将DOCX文件转换为Markdown格式"""
    try:
        file_path = arguments.get("file_path")

        if not file_path:
            return "错误: 需要提供文档文件路径"

        doc_path = Path(file_path)
        if not doc_path.exists():
            return f"错误: 文件不存在: {file_path}"

        if not doc_path.suffix.lower() in ['.docx', '.doc']:
            return "错误: 文件必须是DOC或DOCX格式"

        doc = Document(str(doc_path))
        markdown_content = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            style_name = paragraph.style.name.lower()

            if 'heading 1' in style_name or 'title' in style_name:
                markdown_content.append(f"# {text}")
            elif 'heading 2' in style_name:
                markdown_content.append(f"## {text}")
            elif 'heading 3' in style_name:
                markdown_content.append(f"### {text}")
            elif 'heading 4' in style_name:
                markdown_content.append(f"#### {text}")
            elif 'heading 5' in style_name:
                markdown_content.append(f"##### {text}")
            elif 'heading 6' in style_name:
                markdown_content.append(f"###### {text}")
            else:
                formatted_text = text
                for run in paragraph.runs:
                    if run.bold and run.text in formatted_text:
                        formatted_text = formatted_text.replace(run.text, f"**{run.text}**")
                    elif run.italic and run.text in formatted_text:
                        formatted_text = formatted_text.replace(run.text, f"*{run.text}*")

                markdown_content.append(formatted_text)

        for table in doc.tables:
            markdown_content.append("\n")

            header_row = []
            for cell in table.rows[0].cells:
                header_row.append(cell.text.strip())
            markdown_content.append("| " + " | ".join(header_row) + " |")

            separator = "| " + " | ".join(["---"] * len(header_row)) + " |"
            markdown_content.append(separator)

            for row in table.rows[1:]:
                data_row = []
                for cell in row.cells:
                    data_row.append(cell.text.strip())
                markdown_content.append("| " + " | ".join(data_row) + " |")

            markdown_content.append("\n")

        result = "\n\n".join(markdown_content)
        return f"成功转换为Markdown！\n\n文件: {file_path}\n\n{result}"

    except Exception as e:
        logger.error(f"DOCX转Markdown过程中发生错误: {str(e)}")
        return f"错误: {str(e)}"


async def excel_to_markdown(arguments: Dict[str, Any]) -> str:
    """将Excel文件转换为Markdown格式"""
    try:
        file_path = arguments.get("file_path")
        sheet_name = arguments.get("sheet_name", None)

        if not sheet_name:
            sheet_name = "sheet1"

        if not file_path:
            return "错误: 需要提供Excel文件路径"

        excel_path = Path(file_path)
        if not excel_path.exists():
            return f"错误: 文件不存在: {file_path}"

        if not excel_path.suffix.lower() in ['.xlsx', '.xls']:
            return "错误: 文件必须是Excel格式(.xlsx或.xls)"

        workbook = openpyxl.load_workbook(excel_path, data_only=True)
        markdown_content = []

        if sheet_name:
            if sheet_name in workbook.sheetnames:
                sheets_to_process = [sheet_name]
            else:
                return f"错误: 工作表 '{sheet_name}' 不存在"
        else:
            sheets_to_process = workbook.sheetnames

        for sheet_name in sheets_to_process:
            worksheet = workbook[str(sheet_name)]

            markdown_content.append(f"# {sheet_name}")
            markdown_content.append("")

            if worksheet.max_row == 1 and worksheet.max_column == 1:
                markdown_content.append("*此工作表为空*")
                markdown_content.append("")
                continue

            rows_data = []
            for row in worksheet.iter_rows(min_row=1, max_row=worksheet.max_row,
                                           min_col=1, max_col=worksheet.max_column):
                row_data = []
                for cell in row:
                    value = cell.value
                    if value is None:
                        value = ""
                    row_data.append(str(value))

                if any(cell_value.strip() for cell_value in row_data):
                    rows_data.append(row_data)

            if rows_data:
                header_row = rows_data[0]
                markdown_content.append("| " + " | ".join(header_row) + " |")

                separator = "| " + " | ".join(["---"] * len(header_row)) + " |"
                markdown_content.append(separator)

                for row_data in rows_data[1:]:
                    while len(row_data) < len(header_row):
                        row_data.append("")
                    markdown_content.append("| " + " | ".join(row_data[:len(header_row)]) + " |")

            markdown_content.append("")

        result = "\n".join(markdown_content)
        return f"成功转换为Markdown！\n\n文件: {file_path}\n处理工作表: {', '.join(sheets_to_process)}\n\n{result}"

    except Exception as e:
        logger.error(f"Excel转Markdown过程中发生错误: {str(e)}")
        return f"错误: {str(e)}"


async def pdf_to_text(arguments: Dict[str, Any]) -> str:
    try:
        file_path = arguments.get("file_path")
        page_range = arguments.get("page_range", "all")

        if not file_path:
            return "错误: 需要提供PDF文件路径"

        pdf_path = Path(file_path)
        if not pdf_path.exists():
            return f"错误: 文件不存在: {file_path}"

        if not pdf_path.suffix.lower() == '.pdf':
            return "错误: 文件必须是PDF格式"

        pages_to_extract = await parse_page_range(page_range)

        extracted_text = []
        page_count = 0

        with pdfplumber.open(pdf_path) as pdf:
            total_pages = len(pdf.pages)

            for page_num, page in enumerate(pdf.pages, 1):
                if pages_to_extract and page_num not in pages_to_extract:
                    continue

                page_text = page.extract_text()
                if page_text:
                    extracted_text.append(f"=== 第 {page_num} 页 ===\n{page_text}\n")
                    page_count += 1

        if not extracted_text:
            return "警告: 未能从PDF中提取到任何文本内容"

        full_text = "\n".join(extracted_text)

        return f"成功提取文本！\n\n文件: {file_path}\n总页数: {total_pages}\n提取页数: {page_count}\n\n{full_text}"

    except Exception as e:
        logger.error(f"PDF转文字过程中发生错误: {str(e)}")
        return f"错误: {str(e)}"


async def parse_page_range(page_range: str) -> set[int]:
    if not page_range or page_range.lower() == "all":
        return set()

    pages = set()
    parts = page_range.split(",")

    for part in parts:
        part = part.strip()
        if "-" in part:
            start, end = map(int, part.split("-"))
            pages.update(range(start, end + 1))
        else:
            pages.add(int(part))

    return pages


async def csv_to_markdown(arguments: Dict[str, Any]) -> str:
    """将CSV文件转换为Markdown格式"""
    try:
        file_path = arguments.get("file_path")
        encoding = arguments.get("encoding", "utf-8")
        delimiter = arguments.get("delimiter", ",")

        if not file_path:
            return "错误: 需要提供CSV文件路径"

        csv_path = Path(file_path)
        if not csv_path.exists():
            return f"错误: 文件不存在: {file_path}"

        if not csv_path.suffix.lower() == '.csv':
            return "错误: 文件必须是CSV格式"

        encodings_to_try = [encoding, 'utf-8', 'gbk', 'gb2312', 'latin-1']
        df = None
        used_encoding = None

        for enc in encodings_to_try:
            try:
                df = pd.read_csv(csv_path, encoding=enc, delimiter=delimiter)
                used_encoding = enc
                break
            except UnicodeDecodeError:
                continue
            except Exception as e:
                if enc == encodings_to_try[-1]:
                    raise e
                continue

        if df is None:
            return "错误: 无法读取CSV文件，请检查文件编码"

        df = df.fillna("")

        markdown_content = [f"# CSV文件转换结果", f"**文件**: {file_path}", f"**编码**: {used_encoding}",
                            f"**行数**: {len(df)}", f"**列数**: {len(df.columns)}", ""]

        if len(df) == 0:
            markdown_content.append("*CSV文件为空*")
        else:
            headers = [str(col) for col in df.columns]
            markdown_content.append("| " + " | ".join(headers) + " |")

            separator = "| " + " | ".join(["---"] * len(headers)) + " |"
            markdown_content.append(separator)

            for _, row in df.iterrows():
                row_data = [str(value).replace("|", "\\|").replace("\n", " ") for value in row]
                markdown_content.append("| " + " | ".join(row_data) + " |")

        result = "\n".join(markdown_content)
        return f"成功转换为Markdown！\n\n{result}"

    except Exception as e:
        logger.error(f"CSV转Markdown过程中发生错误: {str(e)}")
        return f"错误: {str(e)}"


@mcp.tool(name="pdf_to_text", description="将PDF文件转换为文本")
async def pdf_convert(file_path: str, page_range="all") -> str:
    """
    file_path: PDF文件的路径
    page_range: 页面范围，例如 '1-5' 或 '1,3,5' (可选)
    """
    result = call_tool("pdf_to_text", {"file_path": file_path, "page_range": page_range})
    return await result


@mcp.tool(name="docx_to_markdown", description="将DOC/DOCX文件转换为Markdown格式")
async def docx_convert(file_path: str) -> str:
    """
    file_path: DOC或DOCX文件的路径
    """
    result = call_tool("docx_to_markdown", {"file_path": file_path})
    return await result


@mcp.tool(name="excel_to_markdown", description="将Excel文件转换为Markdown格式")
async def excel_convert(file_path: str, sheet_name: str = None) -> str:
    """
    file_path: Excel文件的路径
    sheet_name: 指定工作表名称（可选，不指定则转换所有工作表）
    """
    result = call_tool("excel_to_markdown", {"file_path": file_path, "sheet_name": sheet_name})
    return await result


@mcp.tool(name="csv_to_markdown", description="将CSV文件转换为Markdown格式")
async def csv_convert(file_path: str, encoding: str = "utf-8", delimiter: str = ",") -> str:
    """
    file_path: CSV文件的路径
    encoding: 文件编码（可选，默认utf-8）
    delimiter: 分隔符（可选，默认逗号）
    """
    result = call_tool("csv_to_markdown", {"file_path": file_path, "encoding": encoding, "delimiter": delimiter})
    return await result


if __name__ == "__main__":
    mcp.run(transport="stdio")
