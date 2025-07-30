import logging
from pathlib import Path
from docx import Document

logger = logging.getLogger("file_to_markdown")


async def docx_to_markdown(arguments: dict) -> str:
    try:
        file_path = arguments.get("file_path")
        if not file_path:
            return "错误: 需要提供文档文件路径"

        doc_path = Path(file_path)
        if not doc_path.exists():
            return f"错误: 文件不存在: {file_path}"

        if doc_path.suffix.lower() not in ['.doc', '.docx']:
            return "错误: 文件必须是DOC或DOCX格式"

        doc = Document(str(doc_path))
        markdown_content = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            style_name = paragraph.style.name.lower()

            if 'heading 1' in style_name or 'title' in style_name:
                markdown_content.append(f"# {text}")
            elif 'heading 2' in style_name:
                markdown_content.append(f"## {text}")
            elif 'heading 3' in style_name:
                markdown_content.append(f"### {text}")
            elif 'heading 4' in style_name:
                markdown_content.append(f"#### {text}")
            elif 'heading 5' in style_name:
                markdown_content.append(f"##### {text}")
            elif 'heading 6' in style_name:
                markdown_content.append(f"###### {text}")
            else:
                formatted_text = text
                for run in paragraph.runs:
                    if run.bold and run.text in formatted_text:
                        formatted_text = formatted_text.replace(run.text, f"**{run.text}**")
                    elif run.italic and run.text in formatted_text:
                        formatted_text = formatted_text.replace(run.text, f"*{run.text}*")

                markdown_content.append(formatted_text)

        for table in doc.tables:
            markdown_content.append("\n")

            header_row = [cell.text.strip() for cell in table.rows[0].cells]
            markdown_content.append("| " + " | ".join(header_row) + " |")
            markdown_content.append("| " + " | ".join(["---"] * len(header_row)) + " |")

            for row in table.rows[1:]:
                data_row = [cell.text.strip() for cell in row.cells]
                markdown_content.append("| " + " | ".join(data_row) + " |")

        result = "\n\n".join(markdown_content)
        return f"成功转换为Markdown！\n\n文件: {file_path}\n\n{result}"

    except Exception as e:
        logger.error(f"DOCX转Markdown失败: {e}")
        return f"错误: {e}"
